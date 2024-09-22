import os
import re
import pandas as pd # python -m pip install --upgrade pandas
import FreeSimpleGUI as sg # python -m pip install --upgrade FreeSimpleGIU
from datetime import date
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

"""
För att skapa distribuerbart program:
    python -m pip install --upgrade auto_py_to_exe
    python -m auto_py_to_exe
"""

class GenerateFileCover():
    def __init__(self):
        pass

    def main(self):
        self.gui()
        while True:
            event, self.values = self.window.read()
            match event:
                case sg.WIN_CLOSED:
                    break

                case "-CASE_BUTTON_FOLDER_PATH-":
                    self.window["-CASE_TEXT_FOLDER_PATH-"].update(self.values["-CASE_BUTTON_FOLDER_PATH-"])
                
                case "-CASE_GENERATE-":
                    self.window["-OUTPUT-"].update("")
                    if self.check_case_fields():
                        continue
                    print("Genererar aktomslag......")
                    self.generate_base()
                    self.generate_case()

                case "-PERSON_BUTTON_DATA_PATH-":
                    self.window["-OUTPUT-"].update("")
                    self.window["-PERSON_TEXT_DATA_PATH-"].update(self.values["-PERSON_BUTTON_DATA_PATH-"])
                    self.check_person_format()
                
                case "-PERSON_RADIO_NAME_SAME-" | "-PERSON_RADIO_NAME_DIFFERENT-":
                    self.show_person_columns()
                
                case "-PERSON_GENERATE-":
                    self.window["-OUTPUT-"].update("")
                    if self.check_person_input():
                        continue
                    if self.check_person_radio():
                        continue
                    if self.check_person_columns():
                        continue
                    if self.check_person_value():
                        continue
                    self.transform_person_id()
                    self.transform_person_name()
                    print("Genererar aktomslag......")
                    self.generate_base()
                    self.generate_person()
                    self.default_state_person()
        self.window.close()

    # Kontrollerar att samtliga fält är ifyllda.
    def check_case_fields(self):
        bool = False
        if self.values["-CASE_BUTTON_FOLDER_PATH-"] == "":
            self.window["-CASE_BUTTON_FOLDER_PATH-"].update(button_color=self.background_error)
            bool = True
        else:
            self.window["-CASE_BUTTON_FOLDER_PATH-"].update(button_color=sg.theme_button_color())

        for key, value in self.values.items():
            if str(key).startswith("-CASE_FIELD_"):
                if value == "":
                    self.window[key].update(background_color=self.background_error)
                    bool = True
                else:
                    self.window[key].update(background_color=sg.theme_input_background_color())
       
        if bool == True:
            print("Tomma fält har markerats")
        return bool
    
    # Kontrollerar att alla kolunmfält är ifyllda.
    def check_person_columns(self):
        bool = False
        if self.values["-PERSON_LIST_ID-"] == "":
            self.window["-PERSON_LIST_ID-"].update(background_color=self.background_error)
            bool = True
        else:
            self.window["-PERSON_LIST_ID-"].update(background_color=sg.theme_input_background_color())

        if self.values["-PERSON_RADIO_NAME_SAME-"]:
                if self.values["-PERSON_LIST_NAME_SAME-"] == "":
                    self.window["-PERSON_LIST_NAME_SAME-"].update(background_color=self.background_error)
                    self.window["-PERSON_LIST_NAME_FIRST-"].update(background_color=self.background_error)
                    self.window["-PERSON_LIST_NAME_LAST-"].update(background_color=self.background_error)
                    bool = True
                else:
                    self.window["-PERSON_LIST_NAME_SAME-"].update(background_color=sg.theme_input_background_color())
        else:
                if self.values["-PERSON_LIST_NAME_FIRST-"] == "":
                    self.window["-PERSON_LIST_NAME_SAME-"].update(background_color=self.background_error)
                    self.window["-PERSON_LIST_NAME_FIRST-"].update(background_color=self.background_error)
                    bool = True
                else:
                    self.window["-PERSON_LIST_NAME_FIRST-"].update(background_color=sg.theme_input_background_color())

                if self.values["-PERSON_LIST_NAME_LAST-"] == "":
                    self.window["-PERSON_LIST_NAME_SAME-"].update(background_color=self.background_error)
                    self.window["-PERSON_LIST_NAME_LAST-"].update(background_color=self.background_error)
                    bool = True
                else:
                    self.window["-PERSON_LIST_NAME_LAST-"].update(background_color=sg.theme_input_background_color())   
        
        if bool == True:
            print("Tomma fält har markerats")
        return bool            

    # Kontrollerar att datafilen har rätt format, hämtar därefter all individdata.
    def check_person_format(self):
        _, extension = os.path.splitext(self.values["-PERSON_BUTTON_DATA_PATH-"])
        match extension:
            case ".xlsx":
                file = pd.read_excel(self.values["-PERSON_BUTTON_DATA_PATH-"], dtype=object, header=None)   
            case ".csv":
                file = pd.read_csv(self.values["-PERSON_BUTTON_DATA_PATH-"], header=None)
            case _:
                print("Filen som lästes in var varken i .xlsx-, eller .csv-format. Var god välj en annan fil")
                return True
        
        self.rows = file.values.tolist()
        self.window["-PERSON_LIST_ID-"].update(values=[row+1 for row in range(len(self.rows[0]))])
        self.window["-PERSON_LIST_NAME_SAME-"].update(values=[row+1 for row in range(len(self.rows[0]))])
        self.window["-PERSON_LIST_NAME_FIRST-"].update(values=[row+1 for row in range(len(self.rows[0]))])
        self.window["-PERSON_LIST_NAME_LAST-"].update(values=[row+1 for row in range(len(self.rows[0]))])
        self.window["-PERSON_COLUMN_RADIO-"].update(visible=True)  
        print("Filen har lästs in")

    # Kontrollerar att samtliga fält är ifyllda.
    def check_person_input(self):
        bool = False
        if self.values["-PERSON_BUTTON_FOLDER_PATH-"] == "":
            self.window["-PERSON_BUTTON_FOLDER_PATH-"].update(button_color=self.background_error)
            bool = True
        else:
            self.window["-PERSON_BUTTON_FOLDER_PATH-"].update(button_color=sg.theme_button_color())
        
        if self.values["-PERSON_BUTTON_DATA_PATH-"] == "":
            self.window["-PERSON_BUTTON_DATA_PATH-"].update(button_color=self.background_error)
            bool = True
        else:
            self.window["-PERSON_BUTTON_DATA_PATH-"].update(button_color=sg.theme_button_color())
        
        if self.values["-PERSON_FIELD_HANDLINGSSLAG-"] == "":
            self.window["-PERSON_FIELD_HANDLINGSSLAG-"].update(background_color=self.background_error)
            bool = True
        else:
            self.window["-PERSON_FIELD_HANDLINGSSLAG-"].update(background_color=sg.theme_input_background_color())        
        
        if bool == True:
            print("Tomma fält har markerats")
        return bool

    # Kontrollerar att valbara alternativ är ifyllda.
    def check_person_radio(self):
        if self.values["-PERSON_RADIO_CATEGORY_YES-"] == self.values["-PERSON_RADIO_CATEGORY_NO-"] == False:
            print("Ni har glömt att säga ifall filen innehåller en första rad med kategorier eller inte")
            return True
    
        if self.values["-PERSON_RADIO_NAME_SAME-"] == self.values["-PERSON_RADIO_NAME_DIFFERENT-"] == False:
            print("Ni har glömt att säga ifall namn finns i en eller flera kolumner")
            return True

    # Kontrollerar att samma kolumn i datafilen inte har valts flera gånger
    def check_person_value(self):
        #Riktigt lång linje, ursäkta min okunskap i kodning, men betyder: if a == b or a == c or a == d or (c == d if c and d != "")
        if self.values["-PERSON_LIST_ID-"] == self.values["-PERSON_LIST_NAME_SAME-"] or self.values["-PERSON_LIST_ID-"] == self.values["-PERSON_LIST_NAME_FIRST-"] or self.values["-PERSON_LIST_ID-"] == self.values["-PERSON_LIST_NAME_LAST-"] or self.values["-PERSON_LIST_NAME_FIRST-"] == self.values["-PERSON_LIST_NAME_LAST-"] != "":
            self.window["-PERSON_LIST_ID-"].update(background_color=self.background_error)
            self.window["-PERSON_LIST_NAME_SAME-"].update(background_color=self.background_error)
            self.window["-PERSON_LIST_NAME_FIRST-"].update(background_color=self.background_error)
            self.window["-PERSON_LIST_NAME_LAST-"].update(background_color=self.background_error)
            print("Personnummer och eller namn kan inte dela samma kolumn")
            return True

    # Återgång till ursprungsläge.
    def default_state_person(self):
        for key, _ in self.values.items():
            if str(key).startswith("-PERSON_LIST_"):
                self.window[key].update(values=[])
            elif str(key).startswith("-PERSON_RADIO_"):
                self.window[key].reset_group()
        
        self.window["-PERSON_FIELD_HANDLINGSSLAG-"].update("")
        self.window["-PERSON_TEXT_FOLDER_PATH-"].update("")
        self.window["-PERSON_TEXT_DATA_PATH-"].update(".xlsx eller .csv")
        self.window["-PERSON_COLUMN_NAME_SAME-"].update(visible=False)
        self.window["-PERSON_COLUMN_NAME_DIFFERENT-"].update(visible=False)
        self.window["-PERSON_COLUMN_ID-"].update(visible=False)
        self.window["-PERSON_COLUMN_RADIO-"].update(visible=False)
        
    # Skapar grunden till aktomslaget.
    def generate_base(self):
        self.document = Document()

        # Typsnitt
        font = self.document.styles['Normal'].font
        font.name, font.size = "Arial", Pt(12)

        # Orientering och storlek
        section = self.document.sections[-1]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height, section.left_margin = Mm(420), Mm(297), Mm(10)

        # Kolumner, lösning från https://stackoverflow.com/a/46682569 
        section._sectPr.xpath('./w:cols')[0].set(qn('w:num'),'2') # Sätter antalet kolumner till 2

    # Skapar ärendeomslag.
    def generate_case(self):
        # Generera tomma rader
        for _ in range(5):
            paragraph = self.document.add_paragraph(style="Normal")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph.paragraph_format.space_after = Pt(12)
        
        # Generera stående text
        self.document.paragraphs[0].text = self.values["-CASE_FIELD_MYNDIGHET-"]
        self.document.paragraphs[1].text = self.values["-CASE_FIELD_ARKIVBILDARE-"]
        self.document.paragraphs[2].text = self.values["-CASE_FIELD_HANDLINGSSLAG-"]
        self.document.paragraphs[3].text = f"År {self.values['-CASE_FIELD_YEAR-']}"

        # Generera målmapp
        os.makedirs(self.values["-CASE_BUTTON_FOLDER_PATH-"], exist_ok=True)

        # Genera löpnummer och spara aktomslag
        for _ in range(int(self.values["-CASE_FIELD_NUMBER-"]), int(self.values["-CASE_FIELD_NUMBER_LAST-"]) + 1):
            self.document.paragraphs[4].text = f"Akt {self.values['-CASE_FIELD_NUMBER-']}"
            self.document.save(f"{self.values['-CASE_BUTTON_FOLDER_PATH-']}/Aktomslag_{self.values['-CASE_FIELD_YEAR-']}_{self.values['-CASE_FIELD_NUMBER-']}.docx")
            self.values["-CASE_FIELD_NUMBER-"] = int(self.values["-CASE_FIELD_NUMBER-"]) + 1
        print(f"Samtliga aktomslag har genererats för år {self.values['-CASE_FIELD_YEAR-']}")

    # Skapar individomslag
    def generate_person(self):
        # Generera tomma rader
        for _ in range(3):
            paragraph = self.document.add_paragraph(style="Normal")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph.paragraph_format.space_after = Pt(12)

        # Generera målmapp
        os.makedirs(self.values["-PERSON_BUTTON_FOLDER_PATH-"], exist_ok=True)
        
        # Generera stående text
        self.document.paragraphs[0].text = self.values["-PERSON_FIELD_HANDLINGSSLAG-"]

        # Hoppar över första raden ifall det är en kategori
        if self.values["-PERSON_RADIO_CATEGORY_YES-"]:
            self.rows.pop(0)

        # Genera individdata och spara aktomslag
        for row in self.rows:
            self.document.paragraphs[1].text = str(row[int(self.values["-PERSON_LIST_ID-"]) - 1])
            if self.values["-PERSON_RADIO_NAME_SAME-"]:
                self.document.paragraphs[2].text = str(row[int(self.values["-PERSON_LIST_NAME_SAME-"]) - 1])                                               
            else:
                self.document.paragraphs[2].text = f"{row[int(self.values['-PERSON_LIST_NAME_FIRST-']) - 1]} {row[int(self.values['-PERSON_LIST_NAME_LAST-']) - 1]}"                                                    
            self.document.save(f"{self.values['-PERSON_BUTTON_FOLDER_PATH-']}/Aktomslag_{row[int(self.values['-PERSON_LIST_ID-']) - 1]}.docx")
        print("Samtliga aktomslag har genererats")

    # All grafisk information.
    def gui(self):
        sg.theme("Dark Teal 7")
        self.background_error = "#eb9f8a"

        # Fliken för information.
        tab_information = [
            [sg.Text("Detta lilla program är till för att enkelt kunna generera olika typer av aktomslag.")],
            [sg.Text("Först väljer ni den omslagstyp som ni är ute efter via flikarna, fyll därefter i aktuell information i fliken för att generera era aktomslag.")],
            [sg.Text("Fliken Ärendeomslag avser ärenden med löpnummer. Här genereras ett omslag per löpnummer för så många löpnummer som önskas.")],
            [sg.Text("Fliken Individomslag avser akter inom exempelvis socialtjänsten eller skolan. Här genereras ett omslag per individ för så många individer som läses in.")],
            [sg.Text("Vid frågor är ni välkommen att höra av er till linus.gustavsson.1@gmail.com")],
            [sg.Image(f"{os.path.dirname(__file__)}/Graphics/file_cover_image.png", subsample=7)],
            ]

        # Fliken för ärendeomslag 
        tab_case = [
            [sg.Column([
                [sg.FolderBrowse("Välj målmapp", size=(15), target="-CASE_BUTTON_FOLDER_PATH-", key="-CASE_BUTTON_FOLDER_PATH-", pad=3, enable_events=True), sg.Text(key="-CASE_TEXT_FOLDER_PATH-")],
                [sg.Text("Myndighet", size=(15)), sg.Input(key="-CASE_FIELD_MYNDIGHET-"), sg.Text("Exempelvis en kommun, eller en statlig institution",)],
                [sg.Text("Arkivbildare", size=(15)), sg.Input(key="-CASE_FIELD_ARKIVBILDARE-"), sg.Text("Exempelvis Xnämnden eller Xstyrelsen",)],
                [sg.Text("Handlingsslag", size=(15)), sg.Input(key="-CASE_FIELD_HANDLINGSSLAG-"), sg.Text("Exempelvis diarieförda ärenden",)],
                [sg.Text("År", size=(15)), sg.Input(key="-CASE_FIELD_YEAR-", do_not_clear=False), sg.Text("Exempelvis 1987",)],
                [sg.Text("Första numret", size=(15)), sg.Input(key="-CASE_FIELD_NUMBER-", do_not_clear=False), sg.Text("Det första löpnumret som ska genereras")],
                [sg.Text("Sista numret", size=(15)), sg.Input(key="-CASE_FIELD_NUMBER_LAST-", do_not_clear=False), sg.Text("Det sista löpnumret som ska genereras",)],
                [sg.Button("Generera aktomslag", key="-CASE_GENERATE-", size=(30, 2))],            
            ])]
            ]
        
        # Fliken för individomslag.
        tab_person_default = [
            [sg.FolderBrowse("Välj målmapp", size=(15), pad=3, key="-PERSON_BUTTON_FOLDER_PATH-", target="-PERSON_TEXT_FOLDER_PATH-", ), sg.Text(key="-PERSON_TEXT_FOLDER_PATH-", )], 
            [sg.FileBrowse("Välj individdata", size=(15), key="-PERSON_BUTTON_DATA_PATH-", pad=3, target="-PERSON_BUTTON_DATA_PATH-", enable_events=True), sg.Text(".xlsx eller .csv", key="-PERSON_TEXT_DATA_PATH-")],
            [sg.Text("Handlingsslag", size=(15)), sg.Input(key="-PERSON_FIELD_HANDLINGSSLAG-"), sg.Text("Exempelvis personalakt, eller elevhälsovårdsjournal")],
            ]
        
        tab_person_radio = [
            [sg.Text("Börjar individfilen med en rad för kategorier?"), sg.Radio("Ja", key="-PERSON_RADIO_CATEGORY_YES-", group_id=1), sg.Radio("Nej", key="-PERSON_RADIO_CATEGORY_NO-", group_id=1)],
            [sg.Text("Finns för och efternamn i samma-, eller olika kolumner?"), sg.Radio("Samma kolumn", key="-PERSON_RADIO_NAME_SAME-", enable_events=True, group_id=2), sg.Radio("Olika kolumner", key="-PERSON_RADIO_NAME_DIFFERENT-", enable_events=True, group_id=2)],
            ]
        
        tab_person_id = [[sg.Text("Vilken kolumn innehåller personnummer?"), sg.Combo(values=[], key="-PERSON_LIST_ID-", size=(5), readonly=True)]]
        tab_person_column_name_same = [[sg.Text("Vilken kolumn innehåller för- och efternamn?"), sg.Combo(values=[], key="-PERSON_LIST_NAME_SAME-", size=(5), readonly=True)]]
        tab_person_column_name_different = [[sg.Text("Vilken kolumn innehåller förnamn?"), sg.Combo(values=[], key="-PERSON_LIST_NAME_FIRST-", size=(5), readonly=True), sg.Text("Vilken kolumn innehåller efternamn?"), sg.Combo(values=[], key="-PERSON_LIST_NAME_LAST-", size=(5), readonly=True)]]
        tab_person_generate = [[sg.Button("Generera aktomslag", key="-PERSON_GENERATE-", size=(30, 2))]]

        tab_person = [
            [sg.Column(tab_person_default)],
            [sg.Column(tab_person_radio, key="-PERSON_COLUMN_RADIO-", visible=False)],               
            [sg.Column(tab_person_id, key="-PERSON_COLUMN_ID-", visible=False), sg.Column(tab_person_column_name_same, key="-PERSON_COLUMN_NAME_SAME-", visible=False), sg.Column(tab_person_column_name_different, key="-PERSON_COLUMN_NAME_DIFFERENT-", visible=False)],               
            [sg.Column(tab_person_generate)],
            ]       
        
        # Strukturerar alla flikar
        self.layout = [
            [sg.TabGroup([[sg.Tab("Information", tab_information), sg.Tab("Ärendeomslag", tab_case), sg.Tab("Individomslag", tab_person)],])],
            [sg.Output(key="-OUTPUT-", size=(None, 3), expand_x=True)],
            ]

        # Skapar fönstret. 
        self.window = sg.Window(
            title="Generera aktomslag", 
            layout=self.layout, 
            default_element_size=(40, None),
            icon=f"{os.path.dirname(__file__)}/Graphics/titlebar_icon.ico",
            )
    
    # Visar rätt ifyllnadsfält beroende på ifall namn är i samma eller olika kolumner
    def show_person_columns(self):
        self.window["-PERSON_COLUMN_ID-"].update(visible=True)

        if self.values["-PERSON_RADIO_NAME_SAME-"]:
            self.window["-PERSON_COLUMN_NAME_SAME-"].update(visible=True)
            self.window["-PERSON_COLUMN_NAME_DIFFERENT-"].update(visible=False)
        else:
            self.window["-PERSON_COLUMN_NAME_DIFFERENT-"].update(visible=True)
            self.window["-PERSON_COLUMN_NAME_SAME-"].update(visible=False)

    # Försöker ändra formen på personnumret för stämma med YYYY-MM-DD-SSSS
    def transform_person_id(self):
        today = date.today()
        for row in self.rows:
            # Om personnummer är 10-siffror, ex. 9102035555
            if matches := re.search("(\d){10}", str(row[int(self.values["-PERSON_LIST_ID-"]) - 1])):
                # Kontroll ifall födelseåret är senare än aktuellt år, och därmed föregås av 19
                if int(matches[0][:2]) > int(str(today.year)[2:]):
                    row[int(self.values["-PERSON_LIST_ID-"]) - 1] = f"19{matches[0][:2]}-{matches[0][2:4]}-{matches[0][4:6]}-{matches[0][6:]}"
                # Om inte ovan, föregås därmed av 20 
                else:
                    row[int(self.values["-PERSON_LIST_ID-"]) - 1] = f"20{matches[0][:2]}-{matches[0][2:4]}-{matches[0][4:6]}-{matches[0][6:]}"

            # Om personnummer är 12-siffror, ex. 199102035555
            if matches := re.search("(\d){12}", str(row[int(self.values["-PERSON_LIST_ID-"]) - 1])):
                row[int(self.values["-PERSON_LIST_ID-"]) - 1] = f"{matches[0][:4]}-{matches[0][4:6]}-{matches[0][6:8]}-{matches[0][8:]}"
    
    # Försöker ändra fullständigt namn i samma kolumn till ordningen för- och efternamn.
    def transform_person_name(self):
        for row in self.rows:
            try:
                name = str(row[int(self.values["-PERSON_LIST_NAME_SAME-"]) - 1]).split(", ")
                row[int(self.values["-PERSON_LIST_NAME_SAME-"] - 1)] = f"{name[1]} {name[0]}"
            except:
                pass

if __name__ == "__main__":
    generate_file_cover = GenerateFileCover()
    generate_file_cover.main()